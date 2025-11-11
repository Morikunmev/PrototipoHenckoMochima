#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para leer archivos Word (.docx)
"""

from docx import Document

archivo = "TRA3_LorenzoAguila_RichardRocuant_RicardoCayupel.docx"

try:
    doc = Document(archivo)
    
    # Guardar contenido en archivo de texto
    archivo_salida = archivo.replace('.docx', '_contenido.txt')
    
    with open(archivo_salida, 'w', encoding='utf-8') as f:
        for i, paragraph in enumerate(doc.paragraphs):
            texto = paragraph.text
            if texto.strip():  # Solo escribir párrafos no vacíos
                f.write(texto + '\n')
        
        # También extraer texto de tablas
        for table in doc.tables:
            f.write('\n--- TABLA ---\n')
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                if row_text.strip():
                    f.write(row_text + '\n')
            f.write('--- FIN TABLA ---\n\n')
    
    print(f"Contenido extraído y guardado en: {archivo_salida}")
    print(f"Total de párrafos: {len(doc.paragraphs)}")
    print()
    print("PRIMEROS 3000 CARACTERES:")
    print("=" * 80)
    
    # Leer y mostrar primeros caracteres
    with open(archivo_salida, 'r', encoding='utf-8') as f:
        contenido = f.read()
        print(contenido[:3000])
        print()
        print("=" * 80)
        print(f"Total de caracteres: {len(contenido)}")
        
except Exception as e:
    print(f"ERROR: {e}")
    import traceback
    traceback.print_exc()

